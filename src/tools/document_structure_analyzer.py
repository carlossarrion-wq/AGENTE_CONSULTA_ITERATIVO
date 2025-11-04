#!/usr/bin/env python3
"""
Analizador de estructura de documentos grandes.
Extrae la jerarquía de secciones de PDFs y DOCX para navegación progresiva.
"""

import re
from dataclasses import dataclass, asdict
from typing import List, Optional, Tuple, Dict, Any
from pathlib import Path
import logging

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 no disponible. Instalar con: pip3 install PyPDF2")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx no disponible. Instalar con: pip3 install python-docx")


@dataclass
class DocumentSection:
    """Representa una sección del documento"""
    id: str                          # "section_1", "section_1.1", etc.
    title: str                       # "1. Introducción"
    level: int                       # 1, 2, 3 (nivel de jerarquía)
    start_page: int                  # Página de inicio
    end_page: int                    # Página de fin
    start_char: Optional[int] = None # Posición de inicio en el texto completo
    end_char: Optional[int] = None   # Posición de fin en el texto completo
    char_count: Optional[int] = None # Número de caracteres
    parent_id: Optional[str] = None  # ID de la sección padre
    children_ids: List[str] = None   # IDs de subsecciones
    
    def __post_init__(self):
        if self.children_ids is None:
            self.children_ids = []
    
    def to_dict(self) -> Dict[str, Any]:
        """Convierte a diccionario"""
        return asdict(self)


@dataclass
class DocumentStructure:
    """Estructura completa del documento"""
    file_path: str
    file_name: str
    file_type: str                   # "pdf", "docx", "txt"
    total_pages: int
    total_chars: int
    sections: List[DocumentSection]
    extraction_method: str           # "bookmarks", "text_analysis", "headings"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convierte a diccionario"""
        return {
            "file_path": self.file_path,
            "file_name": self.file_name,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "total_chars": self.total_chars,
            "sections": [s.to_dict() for s in self.sections],
            "extraction_method": self.extraction_method,
            "table_of_contents": self.generate_toc()
        }
    
    def generate_toc(self) -> List[Dict[str, Any]]:
        """Genera tabla de contenidos formateada"""
        toc = []
        for section in self.sections:
            toc.append({
                "id": section.id,
                "title": section.title,
                "level": section.level,
                "pages": f"{section.start_page}-{section.end_page}",
                "chars": section.char_count or 0
            })
        return toc
    
    def get_section_by_id(self, section_id: str) -> Optional[DocumentSection]:
        """Obtiene una sección por su ID"""
        for section in self.sections:
            if section.id == section_id:
                return section
        return None


class DocumentStructureAnalyzer:
    """Analizador principal de estructura de documentos"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def analyze(self, file_path: str) -> DocumentStructure:
        """
        Analiza un documento y extrae su estructura.
        
        Args:
            file_path: Ruta al archivo a analizar
            
        Returns:
            DocumentStructure con la jerarquía del documento
            
        Raises:
            ValueError: Si el formato no es soportado
            FileNotFoundError: Si el archivo no existe
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {file_path}")
        
        file_type = path.suffix.lower()
        
        if file_type == '.pdf':
            return self._analyze_pdf(file_path)
        elif file_type in ['.docx', '.doc']:
            return self._analyze_docx(file_path)
        elif file_type in ['.txt', '.md']:
            return self._analyze_text(file_path)
        else:
            raise ValueError(f"Formato no soportado: {file_type}")
    
    def _analyze_pdf(self, file_path: str) -> DocumentStructure:
        """Analiza estructura de un PDF"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 no está instalado. Ejecutar: pip3 install PyPDF2")
        
        path = Path(file_path)
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            total_pages = len(reader.pages)
            
            # Extraer texto completo para contar caracteres
            full_text = ""
            for page in reader.pages:
                full_text += page.extract_text() + "\n"
            
            total_chars = len(full_text)
            
            # Intentar extraer estructura desde bookmarks
            sections = []
            extraction_method = "none"
            
            # Método 1: Usar bookmarks/outlines del PDF
            if reader.outline:
                self.logger.info("Extrayendo estructura desde bookmarks del PDF")
                sections = self._extract_from_pdf_bookmarks(reader, full_text)
                extraction_method = "bookmarks"
            
            # Método 2: Análisis de texto si no hay bookmarks
            if not sections:
                self.logger.info("No hay bookmarks. Analizando texto para detectar secciones")
                sections = self._extract_from_pdf_text(reader, full_text)
                extraction_method = "text_analysis"
            
            # Si no se encontraron secciones, crear una sección por defecto
            if not sections:
                self.logger.warning("No se detectaron secciones. Creando sección única")
                sections = [DocumentSection(
                    id="section_1",
                    title="Documento Completo",
                    level=1,
                    start_page=1,
                    end_page=total_pages,
                    start_char=0,
                    end_char=total_chars,
                    char_count=total_chars
                )]
                extraction_method = "default"
            
            return DocumentStructure(
                file_path=file_path,
                file_name=path.name,
                file_type="pdf",
                total_pages=total_pages,
                total_chars=total_chars,
                sections=sections,
                extraction_method=extraction_method
            )
    
    def _extract_from_pdf_bookmarks(self, reader: PyPDF2.PdfReader, 
                                   full_text: str) -> List[DocumentSection]:
        """Extrae secciones desde los bookmarks del PDF"""
        sections = []
        
        def process_outline(outline, level=1, parent_id=None, section_counter=[0]):
            """Procesa recursivamente los bookmarks"""
            for item in outline:
                if isinstance(item, list):
                    # Es una lista de sub-items
                    process_outline(item, level + 1, parent_id, section_counter)
                else:
                    # Es un bookmark individual
                    section_counter[0] += 1
                    section_id = f"section_{section_counter[0]}"
                    
                    try:
                        # Obtener página del bookmark
                        page_num = reader.get_destination_page_number(item) + 1
                        
                        # Obtener título
                        title = item.title if hasattr(item, 'title') else f"Sección {section_counter[0]}"
                        
                        section = DocumentSection(
                            id=section_id,
                            title=title,
                            level=level,
                            start_page=page_num,
                            end_page=page_num,  # Se actualizará después
                            parent_id=parent_id
                        )
                        
                        sections.append(section)
                        
                        # Procesar hijos si existen
                        if hasattr(item, '/Kids'):
                            process_outline(item['/Kids'], level + 1, section_id, section_counter)
                            
                    except Exception as e:
                        self.logger.warning(f"Error procesando bookmark: {e}")
                        continue
        
        # Procesar outline
        process_outline(reader.outline)
        
        # Actualizar end_page de cada sección
        for i, section in enumerate(sections):
            if i < len(sections) - 1:
                section.end_page = sections[i + 1].start_page - 1
            else:
                section.end_page = len(reader.pages)
        
        return sections
    
    def _extract_from_pdf_text(self, reader: PyPDF2.PdfReader, 
                               full_text: str) -> List[DocumentSection]:
        """Extrae secciones analizando el texto del PDF"""
        sections = []
        
        # Patrones comunes para detectar títulos de secciones
        patterns = [
            # Números con punto: "1. Título", "1.1 Título"
            r'^(\d+(?:\.\d+)*)\s*[.\-:)]?\s+([A-ZÁÉÍÓÚÑ][^\n]{3,100})$',
            # Capítulos: "CAPÍTULO 1", "CHAPTER 1"
            r'^(CAP[ÍI]TULO|CHAPTER)\s+(\d+)[:\s]+([^\n]{3,100})$',
            # Secciones: "SECCIÓN 1", "SECTION 1"
            r'^(SECCI[ÓO]N|SECTION)\s+(\d+)[:\s]+([^\n]{3,100})$',
            # Anexos: "ANEXO A", "APPENDIX A"
            r'^(ANEXO|AP[ÉE]NDICE|APPENDIX)\s+([A-Z\d]+)[:\s]+([^\n]{3,100})$',
        ]
        
        compiled_patterns = [re.compile(p, re.MULTILINE | re.IGNORECASE) for p in patterns]
        
        lines = full_text.split('\n')
        section_counter = 0
        current_char = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                current_char += 1
                continue
            
            # Intentar match con cada patrón
            for pattern in compiled_patterns:
                match = pattern.match(line)
                if match:
                    section_counter += 1
                    section_id = f"section_{section_counter}"
                    
                    # Determinar nivel basado en la numeración
                    if '.' in match.group(1) if len(match.groups()) > 0 else '':
                        level = match.group(1).count('.') + 1
                    else:
                        level = 1
                    
                    # Título es el último grupo capturado
                    title = match.group(-1).strip()
                    
                    section = DocumentSection(
                        id=section_id,
                        title=title,
                        level=level,
                        start_page=1,  # Se calculará mejor después
                        end_page=1,
                        start_char=current_char,
                        end_char=current_char  # Se actualizará después
                    )
                    
                    sections.append(section)
                    break
            
            current_char += len(line) + 1
        
        # Actualizar end_char de cada sección
        for i, section in enumerate(sections):
            if i < len(sections) - 1:
                section.end_char = sections[i + 1].start_char - 1
                section.char_count = section.end_char - section.start_char
            else:
                section.end_char = len(full_text)
                section.char_count = section.end_char - section.start_char
        
        return sections
    
    def _analyze_docx(self, file_path: str) -> DocumentStructure:
        """Analiza estructura de un DOCX"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Ejecutar: pip3 install python-docx")
        
        path = Path(file_path)
        doc = DocxDocument(file_path)
        
        # Extraer texto completo
        full_text = "\n".join([para.text for para in doc.paragraphs])
        total_chars = len(full_text)
        
        # Extraer secciones desde estilos de heading
        sections = []
        section_counter = 0
        current_char = 0
        
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                section_counter += 1
                level = int(para.style.name.replace('Heading', '').strip() or '1')
                
                section = DocumentSection(
                    id=f"section_{section_counter}",
                    title=para.text,
                    level=level,
                    start_page=1,  # DOCX no tiene concepto de páginas fácilmente accesible
                    end_page=1,
                    start_char=current_char,
                    end_char=current_char
                )
                
                sections.append(section)
            
            current_char += len(para.text) + 1
        
        # Actualizar end_char
        for i, section in enumerate(sections):
            if i < len(sections) - 1:
                section.end_char = sections[i + 1].start_char - 1
                section.char_count = section.end_char - section.start_char
            else:
                section.end_char = total_chars
                section.char_count = section.end_char - section.start_char
        
        return DocumentStructure(
            file_path=file_path,
            file_name=path.name,
            file_type="docx",
            total_pages=1,  # No aplicable para DOCX
            total_chars=total_chars,
            sections=sections,
            extraction_method="headings"
        )
    
    def _analyze_text(self, file_path: str) -> DocumentStructure:
        """Analiza estructura de un archivo de texto plano"""
        path = Path(file_path)
        
        with open(file_path, 'r', encoding='utf-8') as f:
            full_text = f.read()
        
        total_chars = len(full_text)
        
        # Para archivos de texto, usar análisis similar al PDF
        sections = self._extract_sections_from_text(full_text)
        
        return DocumentStructure(
            file_path=file_path,
            file_name=path.name,
            file_type="txt",
            total_pages=1,
            total_chars=total_chars,
            sections=sections,
            extraction_method="text_analysis"
        )
    
    def _extract_sections_from_text(self, text: str) -> List[DocumentSection]:
        """Extrae secciones de texto plano"""
        # Similar a _extract_from_pdf_text pero sin el reader
        sections = []
        
        patterns = [
            r'^(\d+(?:\.\d+)*)\s*[.\-:)]?\s+([A-ZÁÉÍÓÚÑ][^\n]{3,100})$',
            r'^(CAP[ÍI]TULO|CHAPTER)\s+(\d+)[:\s]+([^\n]{3,100})$',
            r'^(SECCI[ÓO]N|SECTION)\s+(\d+)[:\s]+([^\n]{3,100})$',
            r'^(ANEXO|AP[ÉE]NDICE|APPENDIX)\s+([A-Z\d]+)[:\s]+([^\n]{3,100})$',
            # Títulos en mayúsculas
            r'^([A-ZÁÉÍÓÚÑ\s]{5,50})$',
        ]
        
        compiled_patterns = [re.compile(p, re.MULTILINE | re.IGNORECASE) for p in patterns]
        
        lines = text.split('\n')
        section_counter = 0
        current_char = 0
        
        for line in lines:
            line = line.strip()
            if not line:
                current_char += 1
                continue
            
            for pattern in compiled_patterns:
                match = pattern.match(line)
                if match:
                    section_counter += 1
                    
                    section = DocumentSection(
                        id=f"section_{section_counter}",
                        title=line,
                        level=1,
                        start_page=1,
                        end_page=1,
                        start_char=current_char,
                        end_char=current_char
                    )
                    
                    sections.append(section)
                    break
            
            current_char += len(line) + 1
        
        # Actualizar end_char
        for i, section in enumerate(sections):
            if i < len(sections) - 1:
                section.end_char = sections[i + 1].start_char - 1
                section.char_count = section.end_char - section.start_char
            else:
                section.end_char = len(text)
                section.char_count = section.end_char - section.start_char
        
        return sections


def main():
    """Función de prueba"""
    import sys
    import json
    
    if len(sys.argv) < 2:
        print("Uso: python3 document_structure_analyzer.py <archivo>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    # Configurar logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    try:
        analyzer = DocumentStructureAnalyzer()
        structure = analyzer.analyze(file_path)
        
        print("\n" + "=" * 80)
        print(f"ESTRUCTURA DEL DOCUMENTO: {structure.file_name}")
        print("=" * 80)
        print(f"Tipo: {structure.file_type}")
        print(f"Páginas: {structure.total_pages}")
        print(f"Caracteres: {structure.total_chars:,}")
        print(f"Método de extracción: {structure.extraction_method}")
        print(f"Secciones encontradas: {len(structure.sections)}")
        print("\n" + "-" * 80)
        print("TABLA DE CONTENIDOS:")
        print("-" * 80)
        
        for section in structure.sections:
            indent = "  " * (section.level - 1)
            print(f"{indent}{section.id}: {section.title}")
            print(f"{indent}  └─ Páginas: {section.start_page}-{section.end_page}, "
                  f"Caracteres: {section.char_count or 0:,}")
        
        # Guardar estructura en JSON
        output_file = f"{file_path}.structure.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(structure.to_dict(), f, indent=2, ensure_ascii=False)
        
        print(f"\n✅ Estructura guardada en: {output_file}")
        
    except Exception as e:
        print(f"❌ Error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
