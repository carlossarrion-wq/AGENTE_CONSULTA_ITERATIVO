import os
import json
import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
import base64
from io import BytesIO
from loguru import logger
from typing import List, Dict, Any, Optional
import hashlib
import xml.etree.ElementTree as ET
from xml.dom import minidom
from docx import Document

# Try to import Wand for EMF/WMF support
try:
    from wand.image import Image as WandImage
    WAND_AVAILABLE = True
    logger.info("EMF/WMF image support enabled via Wand (ImageMagick)")
except ImportError:
    WAND_AVAILABLE = False
    logger.debug("EMF/WMF image support not available (Wand not installed)")

class DocumentLoader:
    def __init__(self, connection_manager, config=None):
        self.conn_manager = connection_manager
        self.config = config or {}
        # AÑADIDO: Soporte para XML y Python
        self.supported_extensions = {'.pdf', '.docx', '.xlsx', '.xls', '.txt', '.png', '.jpg', '.jpeg', '.xml', '.py'}
        
        # Cargar configuración de filtros de imagen
        image_filtering = self.config.get('image_filtering', {})
        self.min_image_width = image_filtering.get('min_width', 500)
        self.min_image_height = image_filtering.get('min_height', 500)
        self.min_image_size = image_filtering.get('min_total_pixels', 200000)
        self.max_images_per_doc = image_filtering.get('max_images_per_document', 50)

    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a single document and extract its content"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()

            if file_extension not in self.supported_extensions:
                logger.warning(f"Unsupported file type: {file_extension}")
                return None

            # Get file metadata
            file_stats = os.stat(file_path)
            file_hash = self._calculate_file_hash(file_path)

            document = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_extension': file_extension,
                'file_size': file_stats.st_size,
                'file_hash': file_hash,
                'content': '',
                'images': [],
                'metadata': {}
            }

            # Extract content based on file type
            if file_extension == '.pdf':
                document = self._load_pdf(file_path, document)
            elif file_extension in ['.xlsx', '.xls']:
                document = self._load_excel(file_path, document)
            elif file_extension == '.txt':
                document = self._load_text(file_path, document)
            elif file_extension in ['.png', '.jpg', '.jpeg']:
                document = self._load_image(file_path, document)
            elif file_extension == '.docx':
                document = self._load_docx(file_path, document)
            elif file_extension == '.xml':  # NUEVO: Soporte XML
                document = self._load_xml(file_path, document)
            elif file_extension == '.py':  # NUEVO: Soporte Python
                document = self._load_python(file_path, document)

            logger.debug(f"Successfully loaded document: {file_path}")
            return document

        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            return None

    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def _load_xml(self, file_path: str, document: Dict) -> Dict:
        """Load XML document and extract structured content"""
        try:
            # Parse XML
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract content in multiple formats
            content_parts = []
            
            # 1. Pretty-printed XML structure
            rough_string = ET.tostring(root, 'unicode')
            reparsed = minidom.parseString(rough_string)
            pretty_xml = reparsed.toprettyxml(indent="  ")
            
            content_parts.append("=== ESTRUCTURA XML ===")
            content_parts.append(pretty_xml)
            
            # 2. Extract all text content
            all_text = self._extract_xml_text(root)
            if all_text.strip():
                content_parts.append("\n=== CONTENIDO TEXTUAL ===")
                content_parts.append(all_text)
            
            # 3. Extract key-value pairs and attributes
            key_values = self._extract_xml_key_values(root)
            if key_values:
                content_parts.append("\n=== DATOS ESTRUCTURADOS ===")
                for key, value in key_values.items():
                    content_parts.append(f"{key}: {value}")
            
            # 4. Create searchable summary
            summary = self._create_xml_summary(root, os.path.basename(file_path))
            content_parts.insert(0, summary)
            
            document['content'] = '\n'.join(content_parts)
            document['metadata'].update({
                'xml_root_tag': root.tag,
                'xml_namespace': root.tag.split('}')[0].strip('{') if '}' in root.tag else None,
                'xml_elements_count': len(list(root.iter())),
                'xml_attributes_count': sum(len(elem.attrib) for elem in root.iter()),
                'has_structured_data': True
            })
            
            return document
            
        except ET.ParseError as e:
            logger.error(f"XML parsing error in {file_path}: {e}")
            # Fallback: treat as text file
            return self._load_text(file_path, document)
        except Exception as e:
            logger.error(f"Error loading XML {file_path}: {e}")
            return document

    def _extract_xml_text(self, element) -> str:
        """Extract all text content from XML element recursively"""
        texts = []
        
        # Get element text
        if element.text and element.text.strip():
            texts.append(element.text.strip())
        
        # Get text from all child elements
        for child in element:
            child_text = self._extract_xml_text(child)
            if child_text:
                texts.append(child_text)
            
            # Get tail text
            if child.tail and child.tail.strip():
                texts.append(child.tail.strip())
        
        return ' '.join(texts)

    def _extract_xml_key_values(self, element, prefix='') -> Dict[str, str]:
        """Extract key-value pairs from XML structure"""
        key_values = {}
        
        # Add attributes
        for attr_name, attr_value in element.attrib.items():
            key = f"{prefix}{element.tag}@{attr_name}" if prefix else f"{element.tag}@{attr_name}"
            key_values[key] = attr_value
        
        # Add element text if it's a leaf node
        if element.text and element.text.strip() and len(list(element)) == 0:
            key = f"{prefix}{element.tag}" if prefix else element.tag
            key_values[key] = element.text.strip()
        
        # Process children
        for child in element:
            child_prefix = f"{prefix}{element.tag}." if prefix else f"{element.tag}."
            child_kvs = self._extract_xml_key_values(child, child_prefix)
            key_values.update(child_kvs)
        
        return key_values

    def _load_python(self, file_path: str, document: Dict) -> Dict:
        """Load Python source code file and extract structured content"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                source_code = file.read()
            
            # Create structured content
            content_parts = []
            
            # 1. Create summary header
            file_name = os.path.basename(file_path)
            summary = self._create_python_summary(source_code, file_name)
            content_parts.append(summary)
            
            # 2. Add the full source code
            content_parts.append("\n=== CÓDIGO FUENTE COMPLETO ===")
            content_parts.append(source_code)
            
            # 3. Extract and document key elements
            code_elements = self._extract_python_elements(source_code)
            
            if code_elements['imports']:
                content_parts.append("\n=== IMPORTS Y DEPENDENCIAS ===")
                content_parts.append('\n'.join(code_elements['imports']))
            
            if code_elements['classes']:
                content_parts.append("\n=== CLASES DEFINIDAS ===")
                for class_info in code_elements['classes']:
                    content_parts.append(f"\nClase: {class_info['name']}")
                    if class_info['docstring']:
                        content_parts.append(f"Documentación: {class_info['docstring']}")
                    if class_info['methods']:
                        content_parts.append(f"Métodos: {', '.join(class_info['methods'])}")
            
            if code_elements['functions']:
                content_parts.append("\n=== FUNCIONES DEFINIDAS ===")
                for func_info in code_elements['functions']:
                    content_parts.append(f"\nFunción: {func_info['name']}")
                    if func_info['docstring']:
                        content_parts.append(f"Documentación: {func_info['docstring']}")
                    if func_info['params']:
                        content_parts.append(f"Parámetros: {', '.join(func_info['params'])}")
            
            if code_elements['constants']:
                content_parts.append("\n=== CONSTANTES Y VARIABLES GLOBALES ===")
                content_parts.append('\n'.join(code_elements['constants']))
            
            document['content'] = '\n'.join(content_parts)
            document['metadata'].update({
                'python_file': True,
                'python_classes': len(code_elements['classes']),
                'python_functions': len(code_elements['functions']),
                'python_imports': len(code_elements['imports']),
                'python_lines': len(source_code.split('\n')),
                'has_structured_code': True,
                'code_language': 'python'
            })
            
            return document
            
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    source_code = file.read()
                document['content'] = f"=== ARCHIVO PYTHON: {os.path.basename(file_path)} ===\n\n{source_code}"
                document['metadata']['python_file'] = True
                document['metadata']['encoding_issue'] = True
                return document
            except Exception as e:
                logger.error(f"Error loading Python file {file_path}: {e}")
                return document
        except Exception as e:
            logger.error(f"Error loading Python file {file_path}: {e}")
            return document

    def _create_python_summary(self, source_code: str, filename: str) -> str:
        """Create a searchable summary of the Python file"""
        lines = source_code.split('\n')
        total_lines = len(lines)
        code_lines = len([l for l in lines if l.strip() and not l.strip().startswith('#')])
        comment_lines = len([l for l in lines if l.strip().startswith('#')])
        
        summary_parts = [
            f"=== ARCHIVO PYTHON: {filename} ===",
            f"Archivo de código fuente Python",
            f"Total de líneas: {total_lines}",
            f"Líneas de código: {code_lines}",
            f"Líneas de comentarios: {comment_lines}",
        ]
        
        # Extract module docstring if present
        if source_code.strip().startswith('"""') or source_code.strip().startswith("'''"):
            docstring_end = source_code.find('"""', 3) if source_code.strip().startswith('"""') else source_code.find("'''", 3)
            if docstring_end > 0:
                module_doc = source_code[3:docstring_end].strip()
                summary_parts.append(f"\nDescripción del módulo:")
                summary_parts.append(module_doc[:300] + "..." if len(module_doc) > 300 else module_doc)
        
        return '\n'.join(summary_parts)

    def _extract_python_elements(self, source_code: str) -> Dict:
        """Extract key elements from Python source code"""
        import re
        
        elements = {
            'imports': [],
            'classes': [],
            'functions': [],
            'constants': []
        }
        
        lines = source_code.split('\n')
        
        # Extract imports
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('import ') or stripped.startswith('from '):
                elements['imports'].append(stripped)
        
        # Extract classes with basic info
        class_pattern = r'^class\s+(\w+)(?:\([^)]*\))?:'
        for i, line in enumerate(lines):
            match = re.match(class_pattern, line.strip())
            if match:
                class_name = match.group(1)
                class_info = {
                    'name': class_name,
                    'docstring': '',
                    'methods': []
                }
                
                # Try to get docstring
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if next_line.startswith('"""') or next_line.startswith("'''"):
                        quote = '"""' if next_line.startswith('"""') else "'''"
                        docstring_lines = [next_line.replace(quote, '')]
                        for j in range(i + 2, min(i + 10, len(lines))):
                            if quote in lines[j]:
                                docstring_lines.append(lines[j].split(quote)[0])
                                break
                            docstring_lines.append(lines[j].strip())
                        class_info['docstring'] = ' '.join(docstring_lines).strip()
                
                # Extract method names
                for j in range(i + 1, min(i + 100, len(lines))):
                    if lines[j].strip().startswith('def '):
                        method_match = re.match(r'def\s+(\w+)', lines[j].strip())
                        if method_match:
                            class_info['methods'].append(method_match.group(1))
                    elif lines[j].strip().startswith('class '):
                        break
                
                elements['classes'].append(class_info)
        
        # Extract top-level functions
        func_pattern = r'^def\s+(\w+)\s*\((.*?)\):'
        in_class = False
        for i, line in enumerate(lines):
            # Track if we're inside a class
            if line.strip().startswith('class '):
                in_class = True
            elif line and not line[0].isspace() and not line.strip().startswith('#'):
                if line.strip().startswith('def '):
                    in_class = False
            
            if not in_class:
                match = re.match(func_pattern, line.strip())
                if match:
                    func_name = match.group(1)
                    params = match.group(2)
                    func_info = {
                        'name': func_name,
                        'params': [p.strip().split('=')[0].strip() for p in params.split(',') if p.strip()],
                        'docstring': ''
                    }
                    
                    # Try to get docstring
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if next_line.startswith('"""') or next_line.startswith("'''"):
                            quote = '"""' if next_line.startswith('"""') else "'''"
                            docstring_lines = [next_line.replace(quote, '')]
                            for j in range(i + 2, min(i + 10, len(lines))):
                                if quote in lines[j]:
                                    docstring_lines.append(lines[j].split(quote)[0])
                                    break
                                docstring_lines.append(lines[j].strip())
                            func_info['docstring'] = ' '.join(docstring_lines).strip()
                    
                    elements['functions'].append(func_info)
        
        # Extract constants (uppercase variables at module level)
        const_pattern = r'^([A-Z_][A-Z0-9_]*)\s*='
        for line in lines:
            match = re.match(const_pattern, line.strip())
            if match:
                elements['constants'].append(line.strip())
        
        return elements

    def _create_xml_summary(self, root, filename) -> str:
        """Create a searchable summary of the XML document"""
        summary_parts = [
            f"=== RESUMEN DEL DOCUMENTO XML: {filename} ===",
            f"Documento XML con elemento raíz: {root.tag}",
        ]
        
        # Add namespace info
        if '}' in root.tag:
            namespace = root.tag.split('}')[0].strip('{')
            summary_parts.append(f"Namespace: {namespace}")
        
        # Count elements by type
        element_counts = {}
        for elem in root.iter():
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            element_counts[tag] = element_counts.get(tag, 0) + 1
        
        if element_counts:
            summary_parts.append("Elementos encontrados:")
            for tag, count in sorted(element_counts.items()):
                summary_parts.append(f"  - {tag}: {count} ocurrencias")
        
        # Add key attributes if any
        all_attrs = set()
        for elem in root.iter():
            all_attrs.update(elem.attrib.keys())
        
        if all_attrs:
            summary_parts.append(f"Atributos utilizados: {', '.join(sorted(all_attrs))}")
        
        # Add content preview
        text_content = self._extract_xml_text(root)
        if text_content:
            preview = text_content[:200] + "..." if len(text_content) > 200 else text_content
            summary_parts.append(f"Vista previa del contenido: {preview}")
        
        return '\n'.join(summary_parts)

    def _load_pdf(self, file_path: str, document: Dict) -> Dict:
        """Load PDF document with text and images"""
        try:
            pdf_document = fitz.open(file_path)
            text_content = []
            images = []
            
            total_images_found = 0
            images_filtered = 0
            images_duplicated = 0
            
            # Track unique images by content hash to avoid duplicates
            seen_image_hashes = set()
            
            # NUEVO: Track images by xref to avoid processing same image multiple times
            # (same image can be referenced on multiple pages in PDF)
            seen_xrefs = set()
            
            # NUEVO: Detailed debug tracking for deep analysis
            debug_image_analysis = []

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)

                # Extract text
                text = page.get_text()
                if text.strip():
                    text_content.append(f"--- Página {page_num + 1} ---\n{text}")

                # Extract images with filtering and deduplication
                image_list = page.get_images()
                total_images_found += len(image_list)
                
                logger.info(f"📄 Page {page_num + 1}: PyMuPDF detected {len(image_list)} image objects")
                
                for img_index, img in enumerate(image_list):
                    # Check if we've reached the maximum
                    if len(images) >= self.max_images_per_doc:
                        logger.info(f"Reached maximum of {self.max_images_per_doc} images for {os.path.basename(file_path)}, skipping remaining images")
                        break
                    
                    try:
                        xref = img[0]
                        
                        # NUEVO: Skip if we've already processed this xref (same image on different pages)
                        if xref in seen_xrefs:
                            logger.debug(f"🔄 Image xref {xref} on page {page_num + 1} already processed on another page, skipping")
                            images_duplicated += 1
                            continue
                        
                        # NUEVO: Get detailed image information from PDF
                        img_info = {
                            'page': page_num + 1,
                            'index': img_index,
                            'xref': xref,
                            'status': 'unknown'
                        }
                        
                        # Try to get image metadata from PDF
                        try:
                            base_image = pdf_document.extract_image(xref)
                            img_info['colorspace'] = base_image.get('colorspace', 'unknown')
                            img_info['ext'] = base_image.get('ext', 'unknown')
                            img_info['smask'] = base_image.get('smask', None)
                            img_info['bpc'] = base_image.get('bpc', 'unknown')  # bits per component
                        except Exception as meta_error:
                            img_info['metadata_error'] = str(meta_error)
                        
                        pix = fitz.Pixmap(pdf_document, xref)
                        
                        # Get image dimensions and properties
                        width = pix.width
                        height = pix.height
                        total_pixels = width * height
                        colorspace = pix.colorspace.name if pix.colorspace else 'unknown'
                        n_components = pix.n
                        has_alpha = pix.alpha
                        
                        img_info['width'] = width
                        img_info['height'] = height
                        img_info['total_pixels'] = total_pixels
                        img_info['colorspace_pix'] = colorspace
                        img_info['n_components'] = n_components
                        img_info['has_alpha'] = has_alpha
                        
                        # Calculate hash BEFORE filtering to track all images
                        try:
                            if pix.n - pix.alpha < 4:  # Can convert to PNG
                                img_data = pix.tobytes("png")
                                img_hash = hashlib.md5(img_data).hexdigest()
                                img_info['hash'] = img_hash[:16]
                                img_info['hash_full'] = img_hash
                            else:
                                img_info['hash'] = 'unconvertible'
                                img_info['hash_full'] = 'unconvertible'
                        except Exception as hash_error:
                            img_info['hash'] = f'error: {hash_error}'
                            img_info['hash_full'] = f'error: {hash_error}'
                        
                        # Check size filters
                        meets_width = width >= self.min_image_width
                        meets_height = height >= self.min_image_height
                        meets_pixels = total_pixels >= self.min_image_size
                        
                        img_info['meets_width_filter'] = meets_width
                        img_info['meets_height_filter'] = meets_height
                        img_info['meets_pixels_filter'] = meets_pixels
                        
                        # Filter by image dimensions
                        if not meets_width or not meets_height or not meets_pixels:
                            img_info['status'] = 'filtered_size'
                            img_info['filter_reason'] = []
                            if not meets_width:
                                img_info['filter_reason'].append(f"width {width} < {self.min_image_width}")
                            if not meets_height:
                                img_info['filter_reason'].append(f"height {height} < {self.min_image_height}")
                            if not meets_pixels:
                                img_info['filter_reason'].append(f"pixels {total_pixels} < {self.min_image_size}")
                            
                            logger.debug(f"🔍 Image {img_index + 1} on page {page_num + 1}: "
                                       f"{width}x{height} ({total_pixels:,} px) - "
                                       f"FILTERED: {', '.join(img_info['filter_reason'])}")
                            images_filtered += 1
                            debug_image_analysis.append(img_info)
                            pix = None
                            continue

                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            # Check for duplicates by content hash (different xref, same content)
                            if img_info['hash_full'] in seen_image_hashes:
                                img_info['status'] = 'duplicate_content'
                                logger.debug(f"🔍 Image {img_index + 1} on page {page_num + 1}: "
                                           f"{width}x{height} - DUPLICATE CONTENT (hash: {img_info['hash']})")
                                images_duplicated += 1
                                debug_image_analysis.append(img_info)
                                pix = None
                                continue
                            
                            # Mark this image as seen (both by xref and content hash)
                            seen_xrefs.add(xref)
                            seen_image_hashes.add(img_info['hash_full'])
                            img_info['status'] = 'extracted'
                            
                            img_base64 = base64.b64encode(img_data).decode()

                            images.append({
                                'page': page_num + 1,
                                'index': img_index,
                                'data': img_base64,
                                'format': 'png',
                                'width': width,
                                'height': height,
                                'size': (width, height),
                                'hash': img_info['hash']
                            })
                            logger.info(f"✅ Image {len(images)} on page {page_num + 1}: "
                                      f"{width}x{height} ({total_pixels:,} px) - "
                                      f"colorspace: {colorspace}, components: {n_components}, "
                                      f"hash: {img_info['hash']}")
                        else:
                            img_info['status'] = 'unsupported_colorspace'
                            img_info['filter_reason'] = f"colorspace components {pix.n - pix.alpha} >= 4"
                            logger.debug(f"🔍 Image {img_index + 1} on page {page_num + 1}: "
                                       f"{width}x{height} - UNSUPPORTED colorspace (n={pix.n}, alpha={pix.alpha})")
                        
                        debug_image_analysis.append(img_info)
                        pix = None
                        
                    except Exception as e:
                        logger.debug(f"Error extracting image {img_index + 1} from page {page_num + 1}: {e}")
                        debug_image_analysis.append({
                            'page': page_num + 1,
                            'index': img_index,
                            'status': 'error',
                            'error': str(e)
                        })
                
                # Break outer loop if max reached
                if len(images) >= self.max_images_per_doc:
                    break

            document['content'] = '\n\n'.join(text_content)
            document['images'] = images
            document['metadata']['total_pages'] = len(pdf_document)
            document['metadata']['total_images'] = len(images)
            document['metadata']['total_images_found'] = total_images_found
            document['metadata']['images_filtered'] = images_filtered
            document['metadata']['images_duplicated'] = images_duplicated
            document['metadata']['debug_image_analysis'] = debug_image_analysis
            
            # NUEVO: Generate detailed analysis report
            if total_images_found > 0:
                logger.info(f"\n{'='*80}")
                logger.info(f"📊 DETAILED IMAGE ANALYSIS: {os.path.basename(file_path)}")
                logger.info(f"{'='*80}")
                logger.info(f"Total image objects detected: {total_images_found}")
                logger.info(f"Filtered (too small): {images_filtered}")
                logger.info(f"Duplicates skipped: {images_duplicated}")
                logger.info(f"Unique images extracted: {len(images)}")
                logger.info(f"Filters applied: {self.min_image_width}x{self.min_image_height} px, {self.min_image_size:,} total pixels")
                
                # Analyze patterns in filtered images
                if images_filtered > 0:
                    logger.info(f"\n📉 FILTERED IMAGES BREAKDOWN:")
                    filtered_images = [img for img in debug_image_analysis if img['status'] == 'filtered_size']
                    
                    # Group by dimensions
                    dimension_groups = {}
                    for img in filtered_images:
                        dim_key = f"{img.get('width', 0)}x{img.get('height', 0)}"
                        if dim_key not in dimension_groups:
                            dimension_groups[dim_key] = []
                        dimension_groups[dim_key].append(img)
                    
                    for dim, imgs in sorted(dimension_groups.items(), key=lambda x: len(x[1]), reverse=True):
                        logger.info(f"  • {dim} px: {len(imgs)} images")
                        if len(imgs) <= 3:
                            for img in imgs:
                                logger.info(f"    - Page {img['page']}, xref {img.get('xref', '?')}, "
                                          f"colorspace: {img.get('colorspace', '?')}")
                
                # Analyze duplicates
                if images_duplicated > 0:
                    logger.info(f"\n🔄 DUPLICATE IMAGES BREAKDOWN:")
                    duplicate_images = [img for img in debug_image_analysis if img['status'] == 'duplicate']
                    
                    # Group by hash
                    hash_groups = {}
                    for img in duplicate_images:
                        hash_key = img.get('hash', 'unknown')
                        if hash_key not in hash_groups:
                            hash_groups[hash_key] = []
                        hash_groups[hash_key].append(img)
                    
                    for hash_val, imgs in sorted(hash_groups.items(), key=lambda x: len(x[1]), reverse=True):
                        pages = [img['page'] for img in imgs]
                        dim = f"{imgs[0].get('width', 0)}x{imgs[0].get('height', 0)}"
                        logger.info(f"  • Hash {hash_val}: {len(imgs)} duplicates, {dim} px, pages: {pages}")
                
                logger.info(f"{'='*80}\n")

            pdf_document.close()
            return document

        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            return document

    def _load_excel(self, file_path: str, document: Dict) -> Dict:
        """Load Excel document"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Convert DataFrame to text representation
                sheet_content = f"--- Hoja: {sheet_name} ---\n"
                sheet_content += df.to_string(index=False)
                content_parts.append(sheet_content)

            document['content'] = '\n\n'.join(content_parts)
            document['metadata']['sheets'] = excel_file.sheet_names
            document['metadata']['total_sheets'] = len(excel_file.sheet_names)

            return document

        except Exception as e:
            logger.error(f"Error loading Excel {file_path}: {e}")
            return document

    def _load_text(self, file_path: str, document: Dict) -> Dict:
        """Load text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                document['content'] = file.read()
            return document
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    document['content'] = file.read()
                return document
            except Exception as e:
                logger.error(f"Error loading text file {file_path}: {e}")
                return document

    def _load_image(self, file_path: str, document: Dict) -> Dict:
        """Load image document with descriptive content"""
        try:
            with open(file_path, 'rb') as file:
                img_data = file.read()
                img_base64 = base64.b64encode(img_data).decode()

                # Get image info
                img = Image.open(file_path)

                # Generate descriptive content for the image
                file_name = os.path.basename(file_path)
                image_format = img.format.upper() if img.format else 'UNKNOWN'
                width, height = img.size

                # Create descriptive content that can be indexed
                descriptive_content = f"""Imagen: {file_name}
Formato: {image_format}
Dimensiones: {width}x{height} píxeles
Tipo de documento: Imagen visual

Esta es una imagen que contiene información visual importante. El archivo se llama "{file_name}" y es de tipo {image_format}.
La imagen tiene unas dimensiones de {width} píxeles de ancho por {height} píxeles de alto.

Contenido visual: Esta imagen puede contener diagramas, gráficos, esquemas, flujos de proceso, interfaces de usuario,
capturas de pantalla, o cualquier otro tipo de información visual relevante para el sistema.

Para análisis detallado del contenido visual, se requiere procesamiento con modelos de visión artificial."""

                document['content'] = descriptive_content
                document['images'] = [{
                    'data': img_base64,
                    'format': img.format.lower() if img.format else 'unknown',
                    'size': img.size
                }]
                document['metadata']['image_format'] = img.format
                document['metadata']['image_size'] = img.size
                document['metadata']['has_visual_content'] = True

            return document

        except Exception as e:
            logger.error(f"Error loading image {file_path}: {e}")
            return document

    def _convert_emf_wmf_to_png(self, image_data: bytes, image_format: str) -> Optional[bytes]:
        """
        Convert EMF/WMF images to PNG using Wand (ImageMagick).
        
        Args:
            image_data: Raw EMF/WMF image bytes
            image_format: Image format string (for logging)
            
        Returns:
            PNG image bytes, or None if conversion fails
        """
        if not WAND_AVAILABLE:
            logger.warning(f"Cannot convert {image_format} image: Wand library not available")
            return None
        
        try:
            # Use Wand to convert EMF/WMF to PNG
            with WandImage(blob=image_data) as img:
                # Convert to PNG format
                img.format = 'png'
                
                # Get the PNG bytes
                png_data = img.make_blob('png')
                
                logger.info(f"Successfully converted {image_format} image to PNG using Wand")
                return png_data
                
        except Exception as e:
            logger.warning(f"Failed to convert {image_format} image using Wand: {e}")
            return None

    def _resize_image_if_needed(self, image_data: bytes, max_pixels: int = 2048 * 2048, image_format: str = 'unknown') -> Optional[bytes]:
        """
        Resize image if it exceeds max pixels to comply with AWS Bedrock limits.
        For EMF/WMF formats, attempts conversion to PNG first.
        
        Args:
            image_data: Raw image bytes
            max_pixels: Maximum allowed pixels (default: 2048x2048 = 4,194,304)
            image_format: Original image format for logging
            
        Returns:
            Resized image bytes in PNG format, or None if format is unsupported
        """
        try:
            # Try to open with PIL first
            img = Image.open(BytesIO(image_data))
            width, height = img.size
            current_pixels = width * height
            
            # Check if resize is needed
            if current_pixels <= max_pixels:
                # Convert to PNG if not already
                if img.format != 'PNG':
                    output = BytesIO()
                    img.save(output, format='PNG')
                    return output.getvalue()
                return image_data
            
            # Calculate new dimensions maintaining aspect ratio
            ratio = (max_pixels / current_pixels) ** 0.5
            new_width = int(width * ratio)
            new_height = int(height * ratio)
            
            # Ensure dimensions don't exceed 2048
            if new_width > 2048:
                new_width = 2048
                new_height = int(height * (2048 / width))
            if new_height > 2048:
                new_height = 2048
                new_width = int(width * (2048 / height))
            
            logger.info(f"Resizing image from {width}x{height} ({current_pixels:,} pixels) to {new_width}x{new_height} ({new_width*new_height:,} pixels)")
            
            # Resize image
            img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # Convert to PNG
            output = BytesIO()
            img_resized.save(output, format='PNG')
            return output.getvalue()
            
        except Exception as e:
            # Check if it's an unsupported format (WMF, EMF, etc.)
            error_str = str(e).lower()
            is_emf_wmf = any(fmt in error_str for fmt in ['wmf', 'emf']) or \
                         any(fmt in image_format.lower() for fmt in ['wmf', 'emf']) or \
                         "cannot find loader" in error_str
            
            if is_emf_wmf:
                logger.info(f"Detected EMF/WMF format ({image_format}), attempting conversion with Wand")
                
                # Try to convert EMF/WMF to PNG using Wand
                png_data = self._convert_emf_wmf_to_png(image_data, image_format)
                
                if png_data:
                    # Now resize the converted PNG if needed
                    return self._resize_image_if_needed(png_data, max_pixels, 'png')
                else:
                    logger.warning(f"Skipping unsupported image format ({image_format}): conversion failed")
                    return None
            else:
                logger.error(f"Error processing image ({image_format}): {e}")
                return image_data

    def _load_docx(self, file_path: str, document: Dict) -> Dict:
        """Load DOCX document and extract text content and embedded images"""
        try:
            # Load the DOCX document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_text = f"\n--- Tabla {table_count} ---\n"
                
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_cells.append(cell_text)
                    
                    if row_cells:
                        table_text += " | ".join(row_cells) + "\n"
                
                if table_text.strip():
                    text_content.append(table_text)
            
            # NUEVO: Extract embedded images from DOCX
            images = []
            image_count = 0
            skipped_count = 0
            converted_count = 0
            
            try:
                # Method 1: Extract images from document relationships
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_part = rel.target_part
                            image_data = image_part.blob
                            
                            # Determine image format from content type
                            content_type = getattr(image_part, 'content_type', 'image/unknown')
                            image_format = content_type.split('/')[-1] if '/' in content_type else 'unknown'
                            
                            # Resize image if needed to comply with Bedrock limits
                            # This will also handle EMF/WMF conversion if Wand is available
                            resized_image_data = self._resize_image_if_needed(image_data, image_format=image_format)
                            
                            # Skip if format is unsupported (returns None)
                            if resized_image_data is None:
                                logger.debug(f"Skipping unsupported image format {image_format} from DOCX relationships")
                                skipped_count += 1
                                continue
                            
                            # Track if this was a converted EMF/WMF
                            if image_format.lower() in ['emf', 'wmf', 'x-emf', 'x-wmf']:
                                converted_count += 1
                            
                            img_base64 = base64.b64encode(resized_image_data).decode()
                            
                            images.append({
                                'data': img_base64,
                                'format': 'png',  # Always PNG after resize
                                'source': 'docx_relationship',
                                'index': image_count,
                                'original_format': image_format
                            })
                            image_count += 1
                            logger.debug(f"Extracted image {image_count} from DOCX relationships: {image_format}")
                            
                        except Exception as img_error:
                            logger.debug(f"Error extracting image from relationship: {img_error}")
                            continue
                
                # Method 2: Extract images from inline shapes in paragraphs
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        # Check for inline shapes (images)
                        if hasattr(run.element, 'xpath'):
                            # Look for drawing elements that contain images
                            try:
                                drawings = run.element.xpath('.//a:blip', namespaces={
                                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                                })
                            except TypeError:
                                # Fallback for older lxml versions that don't support namespaces parameter
                                drawings = run.element.xpath('.//*[local-name()="blip"]')
                            
                            for drawing in drawings:
                                try:
                                    # Get the relationship ID
                                    embed_id = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    if embed_id:
                                        # Get the image from the relationship
                                        image_part = doc.part.related_parts[embed_id]
                                        image_data = image_part.blob
                                        
                                        # Determine image format
                                        content_type = getattr(image_part, 'content_type', 'image/unknown')
                                        image_format = content_type.split('/')[-1] if '/' in content_type else 'unknown'
                                        
                                        # Resize image if needed to comply with Bedrock limits
                                        # This will also handle EMF/WMF conversion if Wand is available
                                        resized_image_data = self._resize_image_if_needed(image_data, image_format=image_format)
                                        
                                        # Skip if format is unsupported (returns None)
                                        if resized_image_data is None:
                                            logger.debug(f"Skipping unsupported image format {image_format} from inline shape")
                                            skipped_count += 1
                                            continue
                                        
                                        # Track if this was a converted EMF/WMF
                                        if image_format.lower() in ['emf', 'wmf', 'x-emf', 'x-wmf']:
                                            converted_count += 1
                                        
                                        img_base64 = base64.b64encode(resized_image_data).decode()
                                        
                                        # Check if we already have this image (avoid duplicates)
                                        image_hash = hashlib.md5(image_data).hexdigest()
                                        existing_hashes = [hashlib.md5(base64.b64decode(img['data'])).hexdigest() 
                                                         for img in images]
                                        
                                        if image_hash not in existing_hashes:
                                            images.append({
                                                'data': img_base64,
                                                'format': 'png',  # Always PNG after resize
                                                'source': 'docx_inline_shape',
                                                'index': image_count,
                                                'hash': image_hash,
                                                'original_format': image_format
                                            })
                                            image_count += 1
                                            logger.debug(f"Extracted inline image {image_count} from DOCX: {image_format}")
                                        
                                except Exception as inline_error:
                                    logger.debug(f"Error extracting inline image: {inline_error}")
                                    continue
                
                # Log summary of image processing
                if converted_count > 0:
                    logger.info(f"Successfully converted {converted_count} EMF/WMF images to PNG")
                if skipped_count > 0:
                    logger.warning(f"Skipped {skipped_count} unsupported images (Wand not available or conversion failed)")
                
            except Exception as image_extraction_error:
                logger.warning(f"Error during image extraction from DOCX {file_path}: {image_extraction_error}")
            
            # Join all content
            full_content = '\n\n'.join(text_content)
            
            # Create descriptive header
            file_name = os.path.basename(file_path)
            header = f"=== DOCUMENTO WORD: {file_name} ===\n"
            
            if full_content.strip():
                document['content'] = header + full_content
            else:
                document['content'] = header + "Documento Word sin contenido textual extraíble."
            
            # Add extracted images to document
            document['images'] = images
            
            # Add metadata
            document['metadata'].update({
                'docx_paragraphs': paragraph_count,
                'docx_tables': table_count,
                'docx_images': len(images),
                'docx_images_converted': converted_count,
                'docx_images_skipped': skipped_count,
                'docx_processed': True,
                'has_structured_content': table_count > 0,
                'has_images': len(images) > 0
            })
            
            logger.debug(f"Successfully extracted DOCX content: {paragraph_count} paragraphs, {table_count} tables, {len(images)} images ({converted_count} converted from EMF/WMF) from {file_path}")
            return document
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            # Fallback to basic implementation
            document['content'] = f"DOCX Document: {os.path.basename(file_path)}\n[Error extracting content: {str(e)}]"
            document['metadata']['docx_error'] = str(e)
            document['images'] = []  # Ensure images list exists even on error
            return document

    def load_documents_from_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Load all supported documents from a directory"""
        documents = []

        if not os.path.exists(directory_path):
            logger.error(f"Directory does not exist: {directory_path}")
            return documents

        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_extension = os.path.splitext(file)[1].lower()

                if file_extension in self.supported_extensions:
                    document = self.load_document(file_path)
                    if document:
                        documents.append(document)

        logger.info(f"Loaded {len(documents)} documents from {directory_path}")
        return documents
